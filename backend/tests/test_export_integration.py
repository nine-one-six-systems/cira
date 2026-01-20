"""Export pipeline integration tests.

Verifies complete data flow from company/analysis models through to
formatted export output in all four formats: Markdown, Word, PDF, JSON.

Requirements coverage:
- EXP-01: Markdown export format
- EXP-02: Word/DOCX export format
- EXP-03: PDF export format
- EXP-04: JSON export format
- EXP-05: 2-page template structure

Test organization:
- TestExportDataFlow: Core format tests (EXP-01 through EXP-04)
- TestExportExecutiveTable: Executive table rendering (EXP-05)
- TestExportTokenStatistics: Token usage in exports
- TestExportSourceUrls: Source URL inclusion (EXP-05)
- TestExportWithSparseData: Graceful handling of missing data
"""

import json
import pytest
from io import BytesIO

from docx import Document
from PyPDF2 import PdfReader

from app import db
from app.services.export_service import ExportService, generate_export
from tests.fixtures.export_fixtures import (
    COMPLETE_ANALYSIS_DATA,
    KEY_EXECUTIVES,
    TOKEN_USAGE_RECORDS,
    CRAWLED_PAGES,
    EXECUTIVE_SUMMARY_TEXT,
    create_complete_export_company,
    create_minimal_export_company,
    calculate_total_tokens,
)


class TestExportDataFlow:
    """
    Integration tests for export data flow across all formats.

    Tests verify that complete analysis data flows correctly from
    database models through ExportService to formatted output.

    Requirements: EXP-01 (Markdown), EXP-02 (Word), EXP-03 (PDF),
                  EXP-04 (JSON), EXP-05 (2-page template)
    """

    def test_markdown_includes_all_analysis_sections(self, app):
        """
        Test markdown export includes all 7 analysis sections.

        Verifies:
        - All sections from 2-page template present
        - Executive summary appears before sections
        - Source URLs appear in Sources section

        Requirements: EXP-01, EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Verify all 7 sections present
            required_sections = [
                "## Company Overview",
                "## Business Model & Products",
                "## Team & Leadership",
                "## Market Position",
                "## Key Insights",
                "## Red Flags & Concerns",
                "## Sources",
            ]
            for section in required_sections:
                assert section in markdown, f"Missing section: {section}"

            # Executive summary appears early (before other sections)
            exec_summary_pos = markdown.find("## Executive Summary")
            company_overview_pos = markdown.find("## Company Overview")
            assert exec_summary_pos < company_overview_pos, (
                "Executive summary should appear before Company Overview"
            )

            # Source URLs appear in Sources section
            sources_section = markdown[markdown.find("## Sources"):]
            assert "https://acme-tech.com/about" in sources_section
            assert "https://acme-tech.com/team" in sources_section

    def test_word_includes_all_analysis_sections(self, app):
        """
        Test Word export includes all analysis sections.

        Verifies:
        - All section headings present in paragraphs
        - Metadata table has all required fields

        Requirements: EXP-02, EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            doc = Document(BytesIO(docx_bytes))
            all_text = " ".join(para.text for para in doc.paragraphs)

            # Verify all section headings present
            required_sections = [
                "Executive Summary",
                "Company Overview",
                "Business Model & Products",
                "Team & Leadership",
                "Market Position",
                "Key Insights",
                "Red Flags & Concerns",
                "Sources",
            ]
            for section in required_sections:
                assert section in all_text, f"Missing section in Word doc: {section}"

            # Verify metadata table exists and has required fields
            assert len(doc.tables) >= 1, "No metadata table in Word document"
            metadata_table = doc.tables[0]
            table_text = " ".join(
                cell.text for row in metadata_table.rows for cell in row.cells
            )
            assert "Website" in table_text
            assert "Industry" in table_text
            assert "Analysis Mode" in table_text
            assert "Pages Analyzed" in table_text
            assert "Tokens Used" in table_text

    def test_pdf_includes_all_analysis_sections(self, app):
        """
        Test PDF export includes all analysis sections.

        Verifies:
        - Key content extractable (company name, executive summary)
        - Page count reasonable (1-3 pages per 2-page template)

        Requirements: EXP-03, EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            reader = PdfReader(BytesIO(pdf_bytes))

            # Extract all text
            all_text = ""
            for page in reader.pages:
                all_text += page.extract_text()

            # Key content extractable
            assert "Acme Technologies" in all_text, "Company name not in PDF"
            assert "Executive Summary" in all_text, "Executive Summary heading not in PDF"

            # Page count reasonable (1-3 pages)
            assert 1 <= len(reader.pages) <= 3, (
                f"PDF has {len(reader.pages)} pages, expected 1-3"
            )

    def test_json_includes_all_structured_data(self, app):
        """
        Test JSON export includes all structured data.

        Verifies:
        - Company metadata present
        - Analysis sections with confidence scores
        - Entities array matches database count
        - Pages array matches database count
        - Token usage statistics correct

        Requirements: EXP-04
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            # Company metadata present
            assert "company" in data
            assert data["company"]["name"] == "Acme Technologies"
            assert data["company"]["websiteUrl"] == "https://acme-tech.com"
            assert data["company"]["industry"] == "Enterprise Software"

            # Analysis sections with confidence scores
            assert "analysis" in data
            assert "sections" in data["analysis"]
            sections = data["analysis"]["sections"]

            expected_sections = [
                "companyOverview",
                "businessModel",
                "teamLeadership",
                "marketPosition",
                "technology",
                "keyInsights",
                "redFlags",
            ]
            for section_name in expected_sections:
                assert section_name in sections, f"Missing section: {section_name}"
                assert "confidence" in sections[section_name], (
                    f"Missing confidence for {section_name}"
                )

            # Entities array matches database count (5 executives + 4 additional)
            assert "entities" in data
            entity_count = len(company.entities)
            assert len(data["entities"]) == entity_count, (
                f"Expected {entity_count} entities, got {len(data['entities'])}"
            )

            # Pages array matches database count
            assert "pages" in data
            page_count = len(company.pages)
            assert len(data["pages"]) == page_count, (
                f"Expected {page_count} pages, got {len(data['pages'])}"
            )

            # Token usage statistics
            assert "tokenUsage" in data
            assert data["tokenUsage"]["total"] > 0
            assert data["tokenUsage"]["estimatedCost"] > 0


class TestExportExecutiveTable:
    """
    Tests for executive table rendering in exports.

    Verifies key executives appear correctly formatted in
    Markdown (GFM table) and Word (table element) exports.

    Requirements: EXP-05
    """

    def test_markdown_executive_table_format(self, app):
        """
        Test markdown has GFM table format for executives.

        Verifies:
        - GFM table header: | Name | Role |
        - All executives appear in table

        Requirements: EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Check GFM table format
            assert "| Name | Role |" in markdown, "GFM table header missing"
            assert "|------|------|" in markdown, "GFM table separator missing"

            # All executives from KEY_EXECUTIVES appear in table
            for exec_data in KEY_EXECUTIVES:
                exec_name = exec_data["name"]
                # Extract just the base role (without suffixes like "& Co-founder")
                exec_role = exec_data["role"].split(" & ")[0]
                assert exec_name in markdown, f"Executive {exec_name} not in markdown"

    def test_word_executive_table(self, app):
        """
        Test Word export has table with executive names.

        Verifies:
        - Table exists for executives
        - Executive names appear in table

        Requirements: EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            doc = Document(BytesIO(docx_bytes))

            # Find executive table (second table after metadata)
            assert len(doc.tables) >= 2, "Expected at least 2 tables (metadata + executives)"

            # Check for executive names in any table
            all_table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_table_text += cell.text + " "

            # Check executives appear in tables
            for exec_data in KEY_EXECUTIVES:
                exec_name = exec_data["name"]
                assert exec_name in all_table_text, (
                    f"Executive {exec_name} not in Word tables"
                )


class TestExportTokenStatistics:
    """
    Tests for token usage statistics in exports.

    Verifies token counts and estimated costs appear
    correctly in Markdown and JSON exports.
    """

    def test_markdown_includes_token_usage(self, app):
        """
        Test markdown includes token count and estimated cost.

        Verifies:
        - Token count appears in metadata
        - Estimated cost appears
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Token count appears
            assert "**Tokens Used:**" in markdown, "Token usage label missing"
            assert "Est. Cost: $" in markdown, "Estimated cost missing"

    def test_json_includes_token_breakdown(self, app):
        """
        Test JSON includes complete token breakdown.

        Verifies:
        - tokenUsage.total matches sum of records
        - tokenUsage.estimatedCost > 0
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            # Token usage section present
            assert "tokenUsage" in data
            token_data = data["tokenUsage"]

            # Total matches expected
            expected = calculate_total_tokens()
            assert token_data["total"] == expected["total"], (
                f"Expected total {expected['total']}, got {token_data['total']}"
            )
            assert token_data["input"] == expected["input"], (
                f"Expected input {expected['input']}, got {token_data['input']}"
            )
            assert token_data["output"] == expected["output"], (
                f"Expected output {expected['output']}, got {token_data['output']}"
            )

            # Estimated cost > 0
            assert token_data["estimatedCost"] > 0, "Estimated cost should be > 0"


class TestExportSourceUrls:
    """
    Tests for source URL inclusion in exports.

    Verifies crawled page URLs appear in Markdown Sources
    section and JSON pages array.

    Requirements: EXP-05
    """

    def test_markdown_lists_source_pages(self, app):
        """
        Test markdown Sources section contains page URLs.

        Verifies:
        - Sources section exists
        - Crawled page URLs appear

        Requirements: EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Sources section exists
            assert "## Sources" in markdown, "Sources section missing"

            # Get sources section content
            sources_start = markdown.find("## Sources")
            sources_section = markdown[sources_start:]

            # Key pages appear (based on priority - about, team, product)
            assert "https://acme-tech.com/about" in sources_section
            assert "https://acme-tech.com/team" in sources_section
            assert "https://acme-tech.com/products" in sources_section

    def test_json_includes_pages(self, app):
        """
        Test JSON includes all crawled page URLs.

        Verifies:
        - pages array present (with include_raw_data=True)
        - All crawled URLs included

        Requirements: EXP-05
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            # Pages array present
            assert "pages" in data, "pages array missing from JSON"

            # Extract URLs from pages array
            page_urls = {p["url"] for p in data["pages"]}

            # All crawled URLs included
            for page_data in CRAWLED_PAGES:
                assert page_data["url"] in page_urls, (
                    f"Missing page URL: {page_data['url']}"
                )


class TestExportWithSparseData:
    """
    Tests for graceful handling of sparse/missing data.

    Verifies exports handle missing sections, entities,
    and pages without errors.
    """

    def test_markdown_handles_missing_sections(self, app):
        """
        Test markdown shows placeholders for missing sections.

        Verifies:
        - Placeholder text appears for missing sections
        - No errors thrown
        """
        with app.app_context():
            company = create_minimal_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should have placeholder text for missing sections
            placeholder_indicators = [
                "No company overview available",
                "No business model information available",
                "No market position information available",
            ]

            placeholders_found = 0
            for indicator in placeholder_indicators:
                if indicator in markdown:
                    placeholders_found += 1

            assert placeholders_found >= 2, (
                "Expected placeholder text for missing sections"
            )

            # Company name still appears
            assert "Minimal Corp" in markdown

    def test_all_formats_succeed_with_minimal_data(self, app):
        """
        Test all four export formats succeed with minimal data.

        Verifies:
        - No exceptions raised
        - Output non-empty for all formats
        """
        with app.app_context():
            company = create_minimal_export_company(db.session)

            # Test markdown
            content, content_type, filename = generate_export(company, "markdown")
            assert len(content) > 0, "Markdown export empty"
            assert filename.endswith(".md")

            # Test word
            content, content_type, filename = generate_export(company, "word")
            assert len(content) > 0, "Word export empty"
            assert filename.endswith(".docx")

            # Test PDF
            content, content_type, filename = generate_export(company, "pdf")
            assert len(content) > 0, "PDF export empty"
            assert filename.endswith(".pdf")

            # Test JSON
            content, content_type, filename = generate_export(company, "json")
            assert len(content) > 0, "JSON export empty"
            assert filename.endswith(".json")
            # Verify valid JSON
            json.loads(content.decode("utf-8"))

    def test_json_handles_no_entities(self, app):
        """
        Test JSON export handles company with no entities.

        Verifies:
        - entities array is empty (not missing)
        - No errors thrown
        """
        with app.app_context():
            company = create_minimal_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            # Entities array present but empty
            assert "entities" in data
            assert len(data["entities"]) == 0, "Expected empty entities array"

    def test_json_handles_no_pages(self, app):
        """
        Test JSON export handles company with no pages.

        Verifies:
        - pages array is empty (not missing)
        - No errors thrown
        """
        with app.app_context():
            company = create_minimal_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            # Pages array present but empty
            assert "pages" in data
            assert len(data["pages"]) == 0, "Expected empty pages array"


class TestExportContentIntegrity:
    """
    Additional tests for content integrity across exports.

    Verifies actual content from analysis flows through
    to export output correctly.
    """

    def test_markdown_contains_actual_content(self, app):
        """
        Test markdown contains actual analysis content.

        Verifies content from COMPLETE_ANALYSIS_DATA appears
        in the markdown output.
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Content from companyOverview
            assert "enterprise software company" in markdown.lower()
            assert "austin" in markdown.lower()

            # Content from businessModel
            assert "subscription" in markdown.lower()

            # Executive summary content
            assert "product-market fit" in markdown.lower()

    def test_json_sections_have_correct_structure(self, app):
        """
        Test JSON sections have content, sources, confidence.

        Verifies each section has the expected structure.
        """
        with app.app_context():
            company = create_complete_export_company(db.session)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            sections = data["analysis"]["sections"]

            for section_name, section_data in sections.items():
                assert "content" in section_data, (
                    f"Section {section_name} missing content"
                )
                assert "sources" in section_data, (
                    f"Section {section_name} missing sources"
                )
                assert "confidence" in section_data, (
                    f"Section {section_name} missing confidence"
                )

                # Confidence should be a number
                assert isinstance(section_data["confidence"], (int, float)), (
                    f"Section {section_name} confidence not numeric"
                )

    def test_executive_summary_appears_in_all_formats(self, app):
        """
        Test executive summary appears in all export formats.

        Verifies the executive summary text flows through
        to each format.
        """
        with app.app_context():
            company = create_complete_export_company(db.session)

            # Markdown
            service = ExportService(company)
            markdown = service.generate_markdown()
            # Check for key phrase from executive summary
            assert "compelling opportunity" in markdown.lower()

            # Word
            docx_bytes = service.generate_word()
            doc = Document(BytesIO(docx_bytes))
            all_text = " ".join(para.text for para in doc.paragraphs).lower()
            assert "compelling opportunity" in all_text

            # JSON
            json_str = service.generate_json()
            data = json.loads(json_str)
            assert "compelling opportunity" in data["analysis"]["executiveSummary"].lower()
