"""Tests for Export Generation Service.

Tests all export formats (Markdown, Word, PDF, JSON) per spec 06-export-formats.md.
"""

import json
import pytest
from io import BytesIO
from datetime import datetime, timezone

from docx import Document
from PyPDF2 import PdfReader

from app import db
from app.models.company import Company, Analysis, Entity, Page, TokenUsage
from app.models.enums import (
    CompanyStatus,
    AnalysisMode,
    EntityType,
    PageType,
    ApiCallType,
)
from app.services.export_service import ExportService, generate_export, CIRA_VERSION


def create_test_company(app) -> str:
    """Create a complete test company with all related data and return its ID."""
    with app.app_context():
        company = Company(
            company_name="Test Company Inc",
            website_url="https://testcompany.com",
            industry="Technology",
            analysis_mode=AnalysisMode.THOROUGH,
            status=CompanyStatus.COMPLETED,
            total_tokens_used=15000,
            estimated_cost=0.0825,
            completed_at=datetime.now(timezone.utc),
        )
        db.session.add(company)
        db.session.flush()
        company_id = company.id

        # Add analysis
        analysis = Analysis(
            company_id=company_id,
            version_number=1,
            executive_summary="Test Company Inc is a leading technology firm specializing in AI solutions. The company has shown strong growth and market positioning.",
            full_analysis={
                "companyOverview": {
                    "content": "Test Company Inc was founded in 2015 and is headquartered in San Francisco.",
                    "sources": ["https://testcompany.com/about"],
                    "confidence": 0.9,
                },
                "businessModel": {
                    "content": "B2B SaaS model with subscription-based pricing.",
                    "sources": ["https://testcompany.com/products"],
                    "confidence": 0.85,
                },
                "teamLeadership": {
                    "content": "Strong leadership team with industry experience.",
                    "sources": ["https://testcompany.com/team"],
                    "confidence": 0.8,
                },
                "marketPosition": {
                    "content": "Positioned in the enterprise AI market segment.",
                    "sources": ["https://testcompany.com/about"],
                    "confidence": 0.75,
                },
                "technology": {
                    "content": "Uses modern tech stack including Python, React, and cloud infrastructure.",
                    "sources": ["https://testcompany.com/careers"],
                    "confidence": 0.7,
                },
                "keyInsights": {
                    "content": "- Strong product-market fit\n- Growing customer base\n- Active hiring",
                    "sources": [],
                    "confidence": 0.8,
                },
                "redFlags": {
                    "content": "No significant concerns identified.",
                    "sources": [],
                    "confidence": 0.9,
                },
            },
        )
        db.session.add(analysis)

        # Add entities
        entities = [
            Entity(
                company_id=company_id,
                entity_type=EntityType.PERSON,
                entity_value="John Smith",
                context_snippet="John Smith serves as CEO of Test Company Inc.",
                source_url="https://testcompany.com/team",
                confidence_score=0.95,
                extra_data={"role": "CEO"},
            ),
            Entity(
                company_id=company_id,
                entity_type=EntityType.PERSON,
                entity_value="Jane Doe",
                context_snippet="Jane Doe is the CTO.",
                source_url="https://testcompany.com/team",
                confidence_score=0.92,
                extra_data={"role": "CTO"},
            ),
            Entity(
                company_id=company_id,
                entity_type=EntityType.PRODUCT,
                entity_value="AI Platform",
                context_snippet="Our flagship AI Platform powers enterprise solutions.",
                source_url="https://testcompany.com/products",
                confidence_score=0.88,
            ),
            Entity(
                company_id=company_id,
                entity_type=EntityType.LOCATION,
                entity_value="San Francisco, CA",
                context_snippet="Headquartered in San Francisco, CA.",
                source_url="https://testcompany.com/about",
                confidence_score=0.97,
            ),
        ]
        for entity in entities:
            db.session.add(entity)

        # Add pages
        pages = [
            Page(
                company_id=company_id,
                url="https://testcompany.com/",
                page_type=PageType.OTHER,
                is_external=False,
            ),
            Page(
                company_id=company_id,
                url="https://testcompany.com/about",
                page_type=PageType.ABOUT,
                is_external=False,
            ),
            Page(
                company_id=company_id,
                url="https://testcompany.com/team",
                page_type=PageType.TEAM,
                is_external=False,
            ),
            Page(
                company_id=company_id,
                url="https://testcompany.com/products",
                page_type=PageType.PRODUCT,
                is_external=False,
            ),
            Page(
                company_id=company_id,
                url="https://testcompany.com/careers",
                page_type=PageType.CAREERS,
                is_external=False,
            ),
        ]
        for page in pages:
            db.session.add(page)

        # Add token usage records
        token_usages = [
            TokenUsage(
                company_id=company_id,
                api_call_type=ApiCallType.EXTRACTION,
                section="overview",
                input_tokens=5000,
                output_tokens=2000,
            ),
            TokenUsage(
                company_id=company_id,
                api_call_type=ApiCallType.ANALYSIS,
                section="business_model",
                input_tokens=4000,
                output_tokens=2000,
            ),
            TokenUsage(
                company_id=company_id,
                api_call_type=ApiCallType.SUMMARIZATION,
                section="summary",
                input_tokens=1500,
                output_tokens=500,
            ),
        ]
        for usage in token_usages:
            db.session.add(usage)

        db.session.commit()
        return company_id


class TestExportService:
    """Tests for ExportService class."""

    def test_export_service_initialization(self, app):
        """Test ExportService initializes correctly with company data."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)

            assert service.company == company
            assert service.analysis is not None
            assert len(service.entities) == 4
            assert len(service.pages) == 5
            assert len(service.token_usages) == 3

    def test_calculate_token_stats(self, app):
        """Test token statistics calculation."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            stats = service._calculate_token_stats()

            assert stats["total"] == 15000  # 5000+4000+1500 + 2000+2000+500
            assert stats["input"] == 10500
            assert stats["output"] == 4500
            assert stats["estimatedCost"] > 0

    def test_get_key_executives(self, app):
        """Test extraction of key executives."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            executives = service._get_key_executives()

            assert len(executives) == 2
            assert any(e["name"] == "John Smith" and e["role"] == "CEO" for e in executives)
            assert any(e["name"] == "Jane Doe" and e["role"] == "CTO" for e in executives)

    def test_get_key_pages(self, app):
        """Test key pages extraction with prioritization."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            key_pages = service._get_key_pages(limit=3)

            assert len(key_pages) == 3
            # About and Team pages should be prioritized
            page_types = [p.page_type.value for p in key_pages]
            assert "about" in page_types
            assert "team" in page_types


class TestMarkdownExport:
    """Tests for Markdown export generation."""

    def test_markdown_export_is_valid_utf8(self, app):
        """Test Markdown output is valid UTF-8."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should be valid UTF-8 string
            assert isinstance(markdown, str)
            encoded = markdown.encode("utf-8")
            assert isinstance(encoded, bytes)

    def test_markdown_export_has_all_sections(self, app):
        """Test Markdown export contains all required sections."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            required_sections = [
                "# Test Company Inc - Intelligence Brief",
                "## Executive Summary",
                "## Company Overview",
                "## Business Model & Products",
                "## Team & Leadership",
                "## Market Position",
                "## Key Insights",
                "## Red Flags & Concerns",
                "## Sources",
            ]

            for section in required_sections:
                assert section in markdown, f"Missing section: {section}"

    def test_markdown_export_has_metadata(self, app):
        """Test Markdown export includes metadata."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            assert "**Analysis Date:**" in markdown
            assert "**Website:** https://testcompany.com" in markdown
            assert "**Industry:** Technology" in markdown
            assert "**Analysis Mode:** thorough" in markdown
            assert "**Pages Analyzed:** 5" in markdown
            assert "**Tokens Used:**" in markdown

    def test_markdown_export_has_executive_table(self, app):
        """Test Markdown export has proper executive table."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Check for GFM table format
            assert "| Name | Role |" in markdown
            assert "|------|------|" in markdown
            assert "| John Smith | CEO |" in markdown
            assert "| Jane Doe | CTO |" in markdown

    def test_markdown_export_has_sources(self, app):
        """Test Markdown export includes source URLs."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            assert "- https://testcompany.com/about" in markdown
            assert "- https://testcompany.com/team" in markdown

    def test_markdown_export_has_footer(self, app):
        """Test Markdown export has CIRA version footer."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            assert f"*Generated by CIRA v{CIRA_VERSION}*" in markdown


class TestWordExport:
    """Tests for Word (.docx) export generation."""

    def test_word_export_opens_without_error(self, app):
        """Test Word export can be opened by python-docx."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            # Should be able to open with python-docx
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None

    def test_word_export_has_title(self, app):
        """Test Word export has company title."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            doc = Document(BytesIO(docx_bytes))

            # Find title in document
            title_found = False
            for para in doc.paragraphs:
                if "Test Company Inc - Intelligence Brief" in para.text:
                    title_found = True
                    break
            assert title_found, "Title not found in Word document"

    def test_word_export_has_metadata_table(self, app):
        """Test Word export has metadata table."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            doc = Document(BytesIO(docx_bytes))

            # Check for tables
            assert len(doc.tables) >= 1, "No tables found in Word document"

            # Check first table has metadata
            first_table = doc.tables[0]
            table_text = " ".join(
                cell.text for row in first_table.rows for cell in row.cells
            )
            assert "Website" in table_text
            assert "Industry" in table_text
            assert "testcompany.com" in table_text

    def test_word_export_has_sections(self, app):
        """Test Word export contains all sections."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            doc = Document(BytesIO(docx_bytes))

            # Get all paragraph text
            all_text = " ".join(para.text for para in doc.paragraphs)

            required_sections = [
                "Executive Summary",
                "Company Overview",
                "Business Model",
                "Team & Leadership",
                "Market Position",
                "Key Insights",
                "Red Flags",
                "Sources",
            ]

            for section in required_sections:
                assert section in all_text, f"Missing section: {section}"


class TestPdfExport:
    """Tests for PDF export generation."""

    def test_pdf_export_generates_without_error(self, app):
        """Test PDF export generates without errors."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0

    def test_pdf_export_is_readable(self, app):
        """Test PDF export can be read by PyPDF2."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            reader = PdfReader(BytesIO(pdf_bytes))
            assert len(reader.pages) > 0

    def test_pdf_export_has_content(self, app):
        """Test PDF export has extractable text content."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            reader = PdfReader(BytesIO(pdf_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text()

            assert "Test Company Inc" in text
            assert "Executive Summary" in text
            assert "CIRA" in text

    def test_pdf_export_page_count_reasonable(self, app):
        """Test PDF export fits within reasonable page count (<=3)."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            reader = PdfReader(BytesIO(pdf_bytes))
            assert len(reader.pages) <= 3, "PDF exceeds 3 pages"


class TestJsonExport:
    """Tests for JSON export generation."""

    def test_json_export_is_valid(self, app):
        """Test JSON export is valid JSON."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()

            # Should parse without errors
            data = json.loads(json_str)
            assert isinstance(data, dict)

    def test_json_export_has_company_data(self, app):
        """Test JSON export contains company data."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            assert "company" in data
            assert data["company"]["name"] == "Test Company Inc"
            assert data["company"]["websiteUrl"] == "https://testcompany.com"
            assert data["company"]["industry"] == "Technology"

    def test_json_export_has_analysis_sections(self, app):
        """Test JSON export contains analysis sections."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            assert "analysis" in data
            assert "executiveSummary" in data["analysis"]
            assert "sections" in data["analysis"]

            sections = data["analysis"]["sections"]
            expected_sections = [
                "companyOverview",
                "businessModel",
                "teamLeadership",
                "marketPosition",
                "technology",
                "keyInsights",
                "redFlags",
            ]
            for section in expected_sections:
                assert section in sections, f"Missing section: {section}"
                assert "content" in sections[section]
                assert "sources" in sections[section]
                assert "confidence" in sections[section]

    def test_json_export_has_token_usage(self, app):
        """Test JSON export contains token usage data."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            assert "tokenUsage" in data
            assert "total" in data["tokenUsage"]
            assert "input" in data["tokenUsage"]
            assert "output" in data["tokenUsage"]
            assert "estimatedCost" in data["tokenUsage"]
            assert data["tokenUsage"]["total"] == 15000

    def test_json_export_has_metadata(self, app):
        """Test JSON export contains metadata."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            assert "metadata" in data
            assert "generatedAt" in data["metadata"]
            assert "ciraVersion" in data["metadata"]
            assert "analysisMode" in data["metadata"]
            assert data["metadata"]["ciraVersion"] == CIRA_VERSION

    def test_json_export_with_entities(self, app):
        """Test JSON export includes entities when requested."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            assert "entities" in data
            assert len(data["entities"]) == 4

            # Check entity structure
            entity = data["entities"][0]
            assert "type" in entity
            assert "value" in entity
            assert "context" in entity
            assert "sourceUrl" in entity
            assert "confidence" in entity

    def test_json_export_with_pages(self, app):
        """Test JSON export includes pages when requested."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)

            assert "pages" in data
            assert len(data["pages"]) == 5

            # Check page structure
            page = data["pages"][0]
            assert "url" in page
            assert "pageType" in page
            assert "crawledAt" in page

    def test_json_export_without_raw_data(self, app):
        """Test JSON export excludes raw data when not requested."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=False)
            data = json.loads(json_str)

            assert "entities" not in data
            assert "pages" not in data

    def test_json_export_dates_iso8601(self, app):
        """Test JSON export uses ISO 8601 date format."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)
            json_str = service.generate_json()
            data = json.loads(json_str)

            # Check generatedAt is ISO 8601
            from datetime import datetime

            generated_at = data["metadata"]["generatedAt"]
            # Should parse without error
            datetime.fromisoformat(generated_at.replace("Z", "+00:00"))


class TestGenerateExport:
    """Tests for generate_export convenience function."""

    def test_generate_markdown_export(self, app):
        """Test generating markdown export via convenience function."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            content, content_type, filename = generate_export(company, "markdown")

            assert content_type == "text/markdown; charset=utf-8"
            assert filename.endswith(".md")
            assert b"# Test Company Inc" in content

    def test_generate_word_export(self, app):
        """Test generating Word export via convenience function."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            content, content_type, filename = generate_export(company, "word")

            assert "openxmlformats" in content_type
            assert filename.endswith(".docx")
            assert isinstance(content, bytes)

    def test_generate_pdf_export(self, app):
        """Test generating PDF export via convenience function."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            content, content_type, filename = generate_export(company, "pdf")

            assert content_type == "application/pdf"
            assert filename.endswith(".pdf")
            assert isinstance(content, bytes)

    def test_generate_json_export(self, app):
        """Test generating JSON export via convenience function."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            content, content_type, filename = generate_export(company, "json")

            assert "application/json" in content_type
            assert filename.endswith(".json")
            # Verify it's valid JSON
            json.loads(content.decode("utf-8"))

    def test_generate_export_invalid_format(self, app):
        """Test invalid format raises ValueError."""
        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            with pytest.raises(ValueError, match="Unsupported export format"):
                generate_export(company, "invalid")

    def test_generate_export_filename_sanitization(self, app):
        """Test filename is sanitized for special characters."""
        with app.app_context():
            company = Company(
                company_name="Test / Company & Inc.",
                website_url="https://test.com",
                status=CompanyStatus.COMPLETED,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            _, _, filename = generate_export(company, "markdown")

            assert "/" not in filename
            assert "\\" not in filename


class TestExportWithNoAnalysis:
    """Tests for export behavior when analysis data is missing."""

    def test_export_with_no_analysis(self, app):
        """Test export handles missing analysis gracefully."""
        with app.app_context():
            company = Company(
                company_name="Empty Company",
                website_url="https://empty.com",
                status=CompanyStatus.COMPLETED,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            assert "Empty Company - Intelligence Brief" in markdown
            assert "No executive summary available." in markdown

    def test_export_with_empty_sections(self, app):
        """Test export handles empty analysis sections."""
        with app.app_context():
            company = Company(
                company_name="Partial Company",
                website_url="https://partial.com",
                status=CompanyStatus.COMPLETED,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary exists",
                full_analysis={},  # Empty sections
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            assert "Summary exists" in markdown
            # Should have placeholder text for missing sections
            assert "No company overview available." in markdown


class TestExportPerformance:
    """Tests for export generation performance."""

    def test_export_generation_performance(self, app):
        """Test export generation completes within time limit (<5 seconds)."""
        import time

        company_id = create_test_company(app)

        with app.app_context():
            company = db.session.get(Company, company_id)
            service = ExportService(company)

            for format_func, name in [
                (service.generate_markdown, "Markdown"),
                (service.generate_word, "Word"),
                (service.generate_pdf, "PDF"),
                (service.generate_json, "JSON"),
            ]:
                start = time.time()
                format_func()
                elapsed = time.time() - start

                assert elapsed < 5.0, f"{name} export took {elapsed:.2f}s (>5s limit)"
