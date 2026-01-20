"""Export generation service for company analysis reports.

Generates company intelligence summaries in multiple formats:
- Markdown (UTF-8)
- Word (.docx) using python-docx
- PDF using ReportLab
- JSON with structured data

All exports follow a consistent 2-page summary structure per spec 06-export-formats.md.
"""

import json
import io
import tempfile
import os
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    HRFlowable,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from app.models.company import Company, Analysis, Entity, Page, TokenUsage
from app.models.enums import EntityType


# Version string for exports
CIRA_VERSION = "1.0.0"


class ExportService:
    """Service for generating export files in various formats."""

    def __init__(self, company: Company, analysis: Analysis | None = None):
        """
        Initialize export service with company data.

        Args:
            company: Company model instance with related data loaded.
            analysis: Optional specific analysis to export. If not provided,
                     uses latest analysis for the company.
        """
        self.company = company
        self.analysis = analysis or self._get_latest_analysis()
        self.entities = list(company.entities)
        self.pages = list(company.pages)
        self.token_usages = list(company.token_usages)

    def _get_latest_analysis(self) -> Analysis | None:
        """Get the most recent analysis for the company."""
        if not self.company.analyses:
            return None
        return max(self.company.analyses, key=lambda a: a.version_number)

    def _calculate_token_stats(self) -> dict[str, Any]:
        """Calculate aggregated token statistics."""
        total_input = sum(t.input_tokens for t in self.token_usages)
        total_output = sum(t.output_tokens for t in self.token_usages)
        total = total_input + total_output
        # Pricing based on Claude Sonnet: $3/$15 per 1M tokens
        estimated_cost = (total_input * 0.000003) + (total_output * 0.000015)

        return {
            "total": total,
            "input": total_input,
            "output": total_output,
            "estimatedCost": round(estimated_cost, 4),
        }

    def _get_analysis_sections(self) -> dict[str, Any]:
        """Extract analysis sections from the analysis."""
        if not self.analysis or not self.analysis.full_analysis:
            return {}
        return self.analysis.full_analysis

    def _format_date(self, dt: datetime | None) -> str:
        """Format datetime for display."""
        if not dt:
            return "Unknown"
        return dt.strftime("%Y-%m-%d %H:%M UTC")

    def _get_key_pages(self, limit: int = 10) -> list[Page]:
        """Get key pages for sources section."""
        # Prioritize important page types
        priority_types = ["about", "team", "product", "service", "contact"]
        sorted_pages = sorted(
            self.pages,
            key=lambda p: (
                priority_types.index(p.page_type.value)
                if p.page_type.value in priority_types
                else 100
            ),
        )
        return sorted_pages[:limit]

    def _get_key_executives(self) -> list[dict[str, str]]:
        """Extract key executives from entities."""
        executives = []
        for entity in self.entities:
            if entity.entity_type == EntityType.PERSON:
                extra = entity.extra_data or {}
                role = extra.get("role", "")
                if role:
                    executives.append({"name": entity.entity_value, "role": role})
        return executives[:10]

    # =========================================================================
    # Markdown Export
    # =========================================================================

    def generate_markdown(self) -> str:
        """
        Generate Markdown export of the analysis.

        Returns:
            UTF-8 encoded markdown string following the 2-page summary template.
        """
        sections = self._get_analysis_sections()
        token_stats = self._calculate_token_stats()

        lines = []

        # Header
        lines.append(f"# {self.company.company_name} - Intelligence Brief\n")
        lines.append(
            f"**Analysis Date:** {self._format_date(self.analysis.created_at if self.analysis else None)}"
        )
        lines.append(f"**Website:** {self.company.website_url}")
        lines.append(f"**Industry:** {self.company.industry or 'Not specified'}")
        lines.append(f"**Analysis Mode:** {self.company.analysis_mode.value}")
        lines.append(f"**Pages Analyzed:** {len(self.pages)}")
        lines.append(
            f"**Tokens Used:** {token_stats['total']:,} (Est. Cost: ${token_stats['estimatedCost']:.4f})"
        )
        lines.append("\n---\n")

        # Executive Summary
        lines.append("## Executive Summary\n")
        if self.analysis and self.analysis.executive_summary:
            lines.append(self.analysis.executive_summary)
        else:
            lines.append("*No executive summary available.*")
        lines.append("\n---\n")

        # Company Overview
        lines.append("## Company Overview\n")
        overview = sections.get("companyOverview", {})
        if isinstance(overview, dict):
            content = overview.get("content", "*No company overview available.*")
        else:
            content = str(overview) if overview else "*No company overview available.*"
        lines.append(content)
        lines.append("\n---\n")

        # Business Model & Products
        lines.append("## Business Model & Products\n")
        business = sections.get("businessModel", {})
        if isinstance(business, dict):
            content = business.get("content", "*No business model information available.*")
        else:
            content = str(business) if business else "*No business model information available.*"
        lines.append(content)
        lines.append("\n---\n")

        # Team & Leadership
        lines.append("## Team & Leadership\n")
        executives = self._get_key_executives()
        if executives:
            lines.append("### Key Executives\n")
            lines.append("| Name | Role |")
            lines.append("|------|------|")
            for exec_info in executives:
                lines.append(f"| {exec_info['name']} | {exec_info['role']} |")
            lines.append("")

        team = sections.get("teamLeadership", {})
        if isinstance(team, dict):
            content = team.get("content", "")
        else:
            content = str(team) if team else ""
        if content:
            lines.append(content)
        elif not executives:
            lines.append("*No team information available.*")
        lines.append("\n---\n")

        # Market Position
        lines.append("## Market Position\n")
        market = sections.get("marketPosition", {})
        if isinstance(market, dict):
            content = market.get("content", "*No market position information available.*")
        else:
            content = str(market) if market else "*No market position information available.*"
        lines.append(content)
        lines.append("\n---\n")

        # Key Insights
        lines.append("## Key Insights\n")
        insights = sections.get("keyInsights", {})
        if isinstance(insights, dict):
            content = insights.get("content", "*No key insights available.*")
        else:
            content = str(insights) if insights else "*No key insights available.*"
        lines.append(content)
        lines.append("\n---\n")

        # Red Flags & Concerns
        lines.append("## Red Flags & Concerns\n")
        red_flags = sections.get("redFlags", {})
        if isinstance(red_flags, dict):
            content = red_flags.get("content", "No significant concerns identified.")
        else:
            content = str(red_flags) if red_flags else "No significant concerns identified."
        lines.append(content)
        lines.append("\n---\n")

        # Sources
        lines.append("## Sources\n")
        lines.append("Key pages analyzed:")
        for page in self._get_key_pages():
            lines.append(f"- {page.url}")
        lines.append("\n---\n")

        # Footer
        lines.append(f"*Generated by CIRA v{CIRA_VERSION}*")

        return "\n".join(lines)

    # =========================================================================
    # Word (.docx) Export
    # =========================================================================

    def generate_word(self) -> bytes:
        """
        Generate Word document export of the analysis.

        Returns:
            Binary content of the .docx file.
        """
        doc = Document()

        # Set up styles
        self._setup_word_styles(doc)

        sections = self._get_analysis_sections()
        token_stats = self._calculate_token_stats()

        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"Intelligence Brief - {self.company.company_name}"
        header_para.style = doc.styles["Header"]

        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Generated by CIRA | Page "

        # Title
        title = doc.add_heading(f"{self.company.company_name} - Intelligence Brief", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        metadata_table = doc.add_table(rows=6, cols=2)
        metadata_table.style = "Table Grid"
        metadata = [
            (
                "Analysis Date",
                self._format_date(self.analysis.created_at if self.analysis else None),
            ),
            ("Website", self.company.website_url),
            ("Industry", self.company.industry or "Not specified"),
            ("Analysis Mode", self.company.analysis_mode.value),
            ("Pages Analyzed", str(len(self.pages))),
            (
                "Tokens Used",
                f"{token_stats['total']:,} (Est. Cost: ${token_stats['estimatedCost']:.4f})",
            ),
        ]
        for i, (label, value) in enumerate(metadata):
            metadata_table.rows[i].cells[0].text = label
            metadata_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        if self.analysis and self.analysis.executive_summary:
            doc.add_paragraph(self.analysis.executive_summary)
        else:
            doc.add_paragraph("No executive summary available.")

        # Company Overview
        doc.add_heading("Company Overview", 1)
        overview = sections.get("companyOverview", {})
        content = (
            overview.get("content", "No company overview available.")
            if isinstance(overview, dict)
            else str(overview) or "No company overview available."
        )
        doc.add_paragraph(content)

        # Business Model & Products
        doc.add_heading("Business Model & Products", 1)
        business = sections.get("businessModel", {})
        content = (
            business.get("content", "No business model information available.")
            if isinstance(business, dict)
            else str(business) or "No business model information available."
        )
        doc.add_paragraph(content)

        # Team & Leadership
        doc.add_heading("Team & Leadership", 1)
        executives = self._get_key_executives()
        if executives:
            doc.add_heading("Key Executives", 2)
            exec_table = doc.add_table(rows=len(executives) + 1, cols=2)
            exec_table.style = "Table Grid"
            exec_table.rows[0].cells[0].text = "Name"
            exec_table.rows[0].cells[1].text = "Role"
            for i, exec_info in enumerate(executives):
                exec_table.rows[i + 1].cells[0].text = exec_info["name"]
                exec_table.rows[i + 1].cells[1].text = exec_info["role"]
            doc.add_paragraph()

        team = sections.get("teamLeadership", {})
        content = (
            team.get("content", "")
            if isinstance(team, dict)
            else str(team) or ""
        )
        if content:
            doc.add_paragraph(content)
        elif not executives:
            doc.add_paragraph("No team information available.")

        # Market Position
        doc.add_heading("Market Position", 1)
        market = sections.get("marketPosition", {})
        content = (
            market.get("content", "No market position information available.")
            if isinstance(market, dict)
            else str(market) or "No market position information available."
        )
        doc.add_paragraph(content)

        # Key Insights
        doc.add_heading("Key Insights", 1)
        insights = sections.get("keyInsights", {})
        content = (
            insights.get("content", "No key insights available.")
            if isinstance(insights, dict)
            else str(insights) or "No key insights available."
        )
        doc.add_paragraph(content)

        # Red Flags
        doc.add_heading("Red Flags & Concerns", 1)
        red_flags = sections.get("redFlags", {})
        content = (
            red_flags.get("content", "No significant concerns identified.")
            if isinstance(red_flags, dict)
            else str(red_flags) or "No significant concerns identified."
        )
        doc.add_paragraph(content)

        # Sources
        doc.add_heading("Sources", 1)
        doc.add_paragraph("Key pages analyzed:")
        for page in self._get_key_pages():
            doc.add_paragraph(page.url, style="List Bullet")

        # Footer note
        doc.add_paragraph()
        footer_note = doc.add_paragraph(f"Generated by CIRA v{CIRA_VERSION}")
        footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_word_styles(self, doc: Document) -> None:
        """Set up custom Word document styles."""
        styles = doc.styles

        # Ensure Header style exists
        if "Header" not in styles:
            header_style = styles.add_style("Header", WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.size = Pt(10)
            header_style.font.italic = True

    # =========================================================================
    # PDF Export
    # =========================================================================

    def generate_pdf(self) -> bytes:
        """
        Generate PDF export of the analysis.

        Returns:
            Binary content of the PDF file.
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=LETTER,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = self._get_pdf_styles()
        story = []

        sections = self._get_analysis_sections()
        token_stats = self._calculate_token_stats()

        # Title
        story.append(
            Paragraph(
                f"{self.company.company_name} - Intelligence Brief", styles["Title"]
            )
        )
        story.append(Spacer(1, 12))

        # Metadata
        metadata_data = [
            [
                "Analysis Date",
                self._format_date(self.analysis.created_at if self.analysis else None),
            ],
            ["Website", self.company.website_url],
            ["Industry", self.company.industry or "Not specified"],
            ["Analysis Mode", self.company.analysis_mode.value],
            ["Pages Analyzed", str(len(self.pages))],
            [
                "Tokens Used",
                f"{token_stats['total']:,} (Est. Cost: ${token_stats['estimatedCost']:.4f})",
            ],
        ]
        metadata_table = Table(metadata_data, colWidths=[2 * inch, 4.5 * inch])
        metadata_table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("TEXTCOLOR", (0, 0), (0, -1), colors.darkblue),
                    ("ALIGN", (0, 0), (0, -1), "RIGHT"),
                    ("ALIGN", (1, 0), (1, -1), "LEFT"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(metadata_table)
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 12))

        # Executive Summary
        story.append(Paragraph("Executive Summary", styles["Heading1"]))
        story.append(Spacer(1, 6))
        summary_text = (
            self.analysis.executive_summary
            if self.analysis and self.analysis.executive_summary
            else "No executive summary available."
        )
        story.append(Paragraph(summary_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Company Overview
        story.append(Paragraph("Company Overview", styles["Heading1"]))
        story.append(Spacer(1, 6))
        overview = sections.get("companyOverview", {})
        content = (
            overview.get("content", "No company overview available.")
            if isinstance(overview, dict)
            else str(overview) or "No company overview available."
        )
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Business Model & Products
        story.append(Paragraph("Business Model & Products", styles["Heading1"]))
        story.append(Spacer(1, 6))
        business = sections.get("businessModel", {})
        content = (
            business.get("content", "No business model information available.")
            if isinstance(business, dict)
            else str(business) or "No business model information available."
        )
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Team & Leadership
        story.append(Paragraph("Team & Leadership", styles["Heading1"]))
        story.append(Spacer(1, 6))
        executives = self._get_key_executives()
        if executives:
            story.append(Paragraph("Key Executives", styles["Heading2"]))
            story.append(Spacer(1, 6))
            exec_data = [["Name", "Role"]]
            for exec_info in executives:
                exec_data.append([exec_info["name"], exec_info["role"]])
            exec_table = Table(exec_data, colWidths=[3 * inch, 3.5 * inch])
            exec_table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(exec_table)
            story.append(Spacer(1, 6))

        team = sections.get("teamLeadership", {})
        content = (
            team.get("content", "")
            if isinstance(team, dict)
            else str(team) or ""
        )
        if content:
            story.append(Paragraph(content, styles["Normal"]))
        elif not executives:
            story.append(Paragraph("No team information available.", styles["Normal"]))
        story.append(Spacer(1, 12))

        # Market Position
        story.append(Paragraph("Market Position", styles["Heading1"]))
        story.append(Spacer(1, 6))
        market = sections.get("marketPosition", {})
        content = (
            market.get("content", "No market position information available.")
            if isinstance(market, dict)
            else str(market) or "No market position information available."
        )
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Key Insights
        story.append(Paragraph("Key Insights", styles["Heading1"]))
        story.append(Spacer(1, 6))
        insights = sections.get("keyInsights", {})
        content = (
            insights.get("content", "No key insights available.")
            if isinstance(insights, dict)
            else str(insights) or "No key insights available."
        )
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Red Flags
        story.append(Paragraph("Red Flags & Concerns", styles["Heading1"]))
        story.append(Spacer(1, 6))
        red_flags = sections.get("redFlags", {})
        content = (
            red_flags.get("content", "No significant concerns identified.")
            if isinstance(red_flags, dict)
            else str(red_flags) or "No significant concerns identified."
        )
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

        # Sources
        story.append(Paragraph("Sources", styles["Heading1"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph("Key pages analyzed:", styles["Normal"]))
        for page in self._get_key_pages():
            story.append(
                Paragraph(f"• {page.url}", styles["Normal"])
            )
        story.append(Spacer(1, 12))

        # Footer
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 6))
        story.append(
            Paragraph(
                f"<i>Generated by CIRA v{CIRA_VERSION}</i>",
                styles["Footer"],
            )
        )

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_pdf_styles(self) -> dict[str, ParagraphStyle]:
        """Get custom PDF styles."""
        base_styles = getSampleStyleSheet()

        return {
            "Title": ParagraphStyle(
                "CustomTitle",
                parent=base_styles["Title"],
                fontSize=18,
                fontName="Helvetica-Bold",
                alignment=TA_CENTER,
                spaceAfter=12,
            ),
            "Heading1": ParagraphStyle(
                "CustomH1",
                parent=base_styles["Heading1"],
                fontSize=14,
                fontName="Helvetica-Bold",
                textColor=colors.darkblue,
                spaceAfter=6,
            ),
            "Heading2": ParagraphStyle(
                "CustomH2",
                parent=base_styles["Heading2"],
                fontSize=12,
                fontName="Helvetica-Bold",
                spaceAfter=6,
            ),
            "Normal": ParagraphStyle(
                "CustomNormal",
                parent=base_styles["Normal"],
                fontSize=10,
                fontName="Helvetica",
                leading=14,
            ),
            "Footer": ParagraphStyle(
                "CustomFooter",
                parent=base_styles["Normal"],
                fontSize=9,
                fontName="Helvetica-Oblique",
                alignment=TA_CENTER,
            ),
        }

    # =========================================================================
    # JSON Export
    # =========================================================================

    def generate_json(self, include_raw_data: bool = True) -> str:
        """
        Generate JSON export with all structured data.

        Args:
            include_raw_data: Include entities and pages in the export.

        Returns:
            JSON string with complete analysis data.
        """
        sections = self._get_analysis_sections()
        token_stats = self._calculate_token_stats()

        data = {
            "company": {
                "id": self.company.id,
                "name": self.company.company_name,
                "websiteUrl": self.company.website_url,
                "industry": self.company.industry,
            },
            "analysis": {
                "versionNumber": self.analysis.version_number if self.analysis else 0,
                "createdAt": (
                    self.analysis.created_at.isoformat()
                    if self.analysis and self.analysis.created_at
                    else None
                ),
                "executiveSummary": (
                    self.analysis.executive_summary if self.analysis else None
                ),
                "sections": self._normalize_sections(sections),
            },
            "tokenUsage": {
                "total": token_stats["total"],
                "input": token_stats["input"],
                "output": token_stats["output"],
                "estimatedCost": token_stats["estimatedCost"],
            },
            "metadata": {
                "generatedAt": datetime.now(timezone.utc).isoformat(),
                "ciraVersion": CIRA_VERSION,
                "analysisMode": self.company.analysis_mode.value,
            },
        }

        if include_raw_data:
            data["entities"] = [
                {
                    "type": entity.entity_type.value,
                    "value": entity.entity_value,
                    "context": entity.context_snippet,
                    "sourceUrl": entity.source_url,
                    "confidence": entity.confidence_score,
                }
                for entity in self.entities
            ]
            data["pages"] = [
                {
                    "url": page.url,
                    "pageType": page.page_type.value,
                    "crawledAt": (
                        page.crawled_at.isoformat() if page.crawled_at else None
                    ),
                }
                for page in self.pages
            ]

        return json.dumps(data, indent=2, ensure_ascii=False)

    def _normalize_sections(self, sections: dict[str, Any]) -> dict[str, Any]:
        """Normalize sections to consistent format for JSON export."""
        normalized = {}
        section_keys = [
            "companyOverview",
            "businessModel",
            "teamLeadership",
            "marketPosition",
            "technology",
            "keyInsights",
            "redFlags",
        ]

        for key in section_keys:
            section = sections.get(key, {})
            if isinstance(section, dict):
                normalized[key] = {
                    "content": section.get("content", ""),
                    "sources": section.get("sources", []),
                    "confidence": section.get("confidence", 0.0),
                }
            else:
                normalized[key] = {
                    "content": str(section) if section else "",
                    "sources": [],
                    "confidence": 0.0,
                }

        return normalized


def generate_export(
    company: Company,
    format: str,
    include_raw_data: bool = True,
    analysis: Analysis | None = None,
) -> tuple[bytes | str, str, str]:
    """
    Generate export in the specified format.

    Args:
        company: Company model instance with related data loaded.
        format: Export format ('markdown', 'word', 'pdf', 'json').
        include_raw_data: Include entities/pages in export (JSON only).
        analysis: Optional specific analysis to export.

    Returns:
        Tuple of (content, content_type, filename).

    Raises:
        ValueError: If format is not supported.
    """
    service = ExportService(company, analysis)
    safe_name = company.company_name.replace(" ", "_").replace("/", "_").replace("\\", "_")[:50]
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d")

    if format == "markdown":
        content = service.generate_markdown()
        return (
            content.encode("utf-8"),
            "text/markdown; charset=utf-8",
            f"{safe_name}_analysis_{timestamp}.md",
        )
    elif format == "word":
        content = service.generate_word()
        return (
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{safe_name}_analysis_{timestamp}.docx",
        )
    elif format == "pdf":
        content = service.generate_pdf()
        return (
            content,
            "application/pdf",
            f"{safe_name}_analysis_{timestamp}.pdf",
        )
    elif format == "json":
        content = service.generate_json(include_raw_data)
        return (
            content.encode("utf-8"),
            "application/json; charset=utf-8",
            f"{safe_name}_analysis_{timestamp}.json",
        )
    else:
        raise ValueError(f"Unsupported export format: {format}")
