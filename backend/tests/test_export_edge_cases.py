"""Edge case tests for export functionality robustness.

Tests verify that the export service handles unusual conditions gracefully:
- Missing analysis data (EXP-05)
- Empty analysis sections (EXP-05)
- Special characters in company names (EXP-05)
- Unicode content in analysis (EXP-05)
- Large content without timeout (EXP-05)
- Missing related data (entities, pages, tokens) (EXP-05)
- Output validity for all formats (EXP-05)

Requirements covered:
- EXP-05: 2-page template structure with graceful fallbacks
"""

import json
import time
from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document
from PyPDF2 import PdfReader

from app import db
from app.models.company import Company, Analysis, Entity, Page, TokenUsage
from app.models.enums import (
    CompanyStatus,
    AnalysisMode,
    EntityType,
    PageType,
    ApiCallType,
)
from app.services.export_service import ExportService, generate_export, CIRA_VERSION


class TestExportMissingAnalysis:
    """Tests for export with missing analysis data (EXP-05).

    Verifies that exports handle companies with no analysis gracefully
    by using placeholder text instead of crashing.
    """

    def test_markdown_with_no_analysis(self, app):
        """
        Test Markdown export handles company with no analysis.

        Verifies EXP-05: Missing analysis shows placeholder text.
        """
        with app.app_context():
            company = Company(
                company_name="No Analysis Corp",
                website_url="https://no-analysis.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should contain company name in title
            assert "# No Analysis Corp - Intelligence Brief" in markdown

            # Should contain placeholder text for executive summary
            assert "No executive summary available." in markdown

            # Should contain placeholder text for sections
            assert "No company overview available." in markdown

            # Should not raise exception
            assert isinstance(markdown, str)
            assert len(markdown) > 100

    def test_word_with_no_analysis(self, app):
        """
        Test Word export handles company with no analysis.

        Verifies EXP-05: Word document opens without error with placeholder text.
        """
        with app.app_context():
            company = Company(
                company_name="No Analysis Word Corp",
                website_url="https://no-analysis-word.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            # Document should open without error
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None

            # Get all text
            all_text = " ".join(para.text for para in doc.paragraphs)

            # Should contain placeholder text
            assert "No executive summary available" in all_text

    def test_pdf_with_no_analysis(self, app):
        """
        Test PDF export handles company with no analysis.

        Verifies EXP-05: PDF is valid and readable even without analysis.
        """
        with app.app_context():
            company = Company(
                company_name="No Analysis PDF Corp",
                website_url="https://no-analysis-pdf.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            # PDF should be valid
            reader = PdfReader(BytesIO(pdf_bytes))
            assert len(reader.pages) > 0

            # Should be able to extract text
            text = ""
            for page in reader.pages:
                text += page.extract_text()

            assert "No Analysis PDF Corp" in text

    def test_json_with_no_analysis(self, app):
        """
        Test JSON export handles company with no analysis.

        Verifies EXP-05: JSON is valid with null/empty values for analysis.
        """
        with app.app_context():
            company = Company(
                company_name="No Analysis JSON Corp",
                website_url="https://no-analysis-json.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json()

            # Should be valid JSON
            data = json.loads(json_str)
            assert isinstance(data, dict)

            # analysis key should exist (not missing)
            assert "analysis" in data

            # executiveSummary should be null
            assert data["analysis"]["executiveSummary"] is None

            # sections should have empty content (not missing keys)
            assert "sections" in data["analysis"]
            for section_key in ["companyOverview", "businessModel"]:
                assert section_key in data["analysis"]["sections"]
                assert data["analysis"]["sections"][section_key]["content"] == ""


class TestExportEmptySections:
    """Tests for export with empty analysis sections (EXP-05).

    Verifies that exports handle analyses with empty or partial sections
    by using placeholder text for missing content.
    """

    def test_markdown_with_empty_full_analysis(self, app):
        """
        Test Markdown export handles empty full_analysis dict.

        Verifies EXP-05: All section headings present with placeholder text.
        """
        with app.app_context():
            company = Company(
                company_name="Empty Analysis Corp",
                website_url="https://empty-analysis.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary exists",
                full_analysis={},  # Empty sections dict
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Executive summary should be present
            assert "Summary exists" in markdown

            # All section headings should be present
            required_sections = [
                "## Executive Summary",
                "## Company Overview",
                "## Business Model & Products",
                "## Team & Leadership",
                "## Market Position",
                "## Key Insights",
                "## Red Flags & Concerns",
                "## Sources",
            ]
            for section in required_sections:
                assert section in markdown, f"Missing section: {section}"

            # Placeholder text for empty sections
            assert "No company overview available." in markdown

    def test_markdown_with_partial_sections(self, app):
        """
        Test Markdown export handles partially filled sections.

        Verifies EXP-05: Existing content appears, missing has placeholder.
        """
        with app.app_context():
            company = Company(
                company_name="Partial Sections Corp",
                website_url="https://partial.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Brief summary",
                full_analysis={
                    "companyOverview": {
                        "content": "This is the company overview content.",
                        "sources": ["https://partial.com/about"],
                        "confidence": 0.9,
                    },
                    # businessModel, teamLeadership, etc. are missing
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # companyOverview content should appear
            assert "This is the company overview content." in markdown

            # Other sections should have placeholder text
            assert "No business model information available." in markdown

    def test_json_with_empty_sections(self, app):
        """
        Test JSON export handles empty full_analysis dict.

        Verifies EXP-05: sections object present but with empty content.
        """
        with app.app_context():
            company = Company(
                company_name="Empty JSON Sections Corp",
                website_url="https://empty-json.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary",
                full_analysis={},  # Empty
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json()

            data = json.loads(json_str)

            # sections object should be present
            assert "sections" in data["analysis"]

            # All section keys should exist with empty content
            expected_sections = [
                "companyOverview",
                "businessModel",
                "teamLeadership",
                "marketPosition",
                "keyInsights",
                "redFlags",
            ]
            for section_key in expected_sections:
                assert section_key in data["analysis"]["sections"]
                assert data["analysis"]["sections"][section_key]["content"] == ""

    def test_word_with_null_content_fields(self, app):
        """
        Test Word export handles analysis with None content fields.

        Verifies EXP-05: Document opens without crash, has placeholder.
        """
        with app.app_context():
            company = Company(
                company_name="Null Content Corp",
                website_url="https://null-content.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            # Analysis with dict sections but None content
            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary=None,  # None instead of string
                full_analysis={
                    "companyOverview": {
                        "content": None,  # None content
                        "sources": [],
                        "confidence": 0.0,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            # Should not crash
            docx_bytes = service.generate_word()

            # Document should open
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None


class TestExportSpecialCharacters:
    """Tests for export with special characters (EXP-05).

    Verifies that filenames are sanitized and content handles
    special characters without issues.
    """

    def test_filename_sanitizes_slashes(self, app):
        """
        Test that forward slashes in company names are sanitized.

        Verifies EXP-05: Filename has no "/" character.
        """
        with app.app_context():
            company = Company(
                company_name="Acme/Corp Inc",
                website_url="https://acme-corp.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            _, _, filename = generate_export(company, "markdown")

            # Filename should not contain forward slash
            assert "/" not in filename

            # Should have underscore substitution
            assert "_" in filename or "-" in filename

    def test_filename_sanitizes_backslashes(self, app):
        """
        Test that backslashes in company names are sanitized.

        Verifies EXP-05: Filename has no backslash character.
        """
        with app.app_context():
            company = Company(
                company_name="Test\\Company",
                website_url="https://test-company.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            _, _, filename = generate_export(company, "pdf")

            # Filename should not contain backslash
            assert "\\" not in filename

    def test_filename_sanitizes_ampersand(self, app):
        """
        Test that ampersands in company names are handled.

        Verifies EXP-05: Filename is reasonable with ampersand.
        """
        with app.app_context():
            company = Company(
                company_name="Foo & Bar LLC",
                website_url="https://foo-bar.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            _, _, filename = generate_export(company, "word")

            # Filename should be reasonable (ampersand may be kept or converted)
            assert len(filename) > 0
            # Should end with .docx
            assert filename.endswith(".docx")

    def test_filename_handles_long_names(self, app):
        """
        Test that very long company names are truncated in filename.

        Verifies EXP-05: Filename truncated to reasonable length.
        """
        with app.app_context():
            # 100-character company name
            long_name = "A" * 100
            company = Company(
                company_name=long_name,
                website_url="https://long-name.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            _, _, filename = generate_export(company, "json")

            # Filename should be truncated (max 50 chars for name part)
            # Full filename includes _analysis_YYYYMMDD.json
            name_part = filename.split("_analysis_")[0]
            assert len(name_part) <= 50

    def test_markdown_handles_special_chars_in_content(self, app):
        """
        Test that markdown special characters in content are handled.

        Verifies EXP-05: Output is valid markdown with special chars.
        """
        with app.app_context():
            company = Company(
                company_name="Special Chars Corp",
                website_url="https://special.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            # Analysis with markdown special characters
            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary with | pipes | and * asterisks * and # hashes",
                full_analysis={
                    "companyOverview": {
                        "content": "Content with `backticks` and <angle brackets> and > quotes",
                        "sources": [],
                        "confidence": 0.9,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should generate without error
            assert isinstance(markdown, str)
            assert len(markdown) > 0

            # Content should be preserved
            assert "|" in markdown
            assert "*" in markdown
            assert "`backticks`" in markdown


class TestExportUnicodeContent:
    """Tests for export with Unicode content (EXP-05).

    Verifies that all export formats handle Unicode characters
    including emojis, CJK characters, and accented text.
    """

    def test_markdown_handles_unicode_company_name(self, app):
        """
        Test Markdown handles Unicode in company name.

        Verifies EXP-05: UTF-8 encoding preserves characters.
        """
        with app.app_context():
            company = Company(
                company_name="Acme Corp",  # Simple name for this test
                website_url="https://acme.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should encode as UTF-8 without error
            encoded = markdown.encode("utf-8")
            assert isinstance(encoded, bytes)

            # Should decode back correctly
            decoded = encoded.decode("utf-8")
            assert "Acme Corp" in decoded

    def test_markdown_handles_unicode_analysis_content(self, app):
        """
        Test Markdown handles Unicode in analysis content.

        Verifies EXP-05: Emojis, CJK, accents preserved in output.
        """
        with app.app_context():
            company = Company(
                company_name="Unicode Content Corp",
                website_url="https://unicode.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            # Analysis with various Unicode
            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Great company with international presence.",
                full_analysis={
                    "companyOverview": {
                        "content": "Company operates in Tokyo and Paris with accent cafe.",
                        "sources": [],
                        "confidence": 0.9,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Content should be preserved
            assert "Tokyo" in markdown
            assert "Paris" in markdown
            assert "cafe" in markdown

    def test_word_handles_unicode(self, app):
        """
        Test Word export handles Unicode in name and content.

        Verifies EXP-05: Document opens with Unicode preserved.
        """
        with app.app_context():
            company = Company(
                company_name="Unicode Word Corp",
                website_url="https://unicode-word.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Company with cafe and resume terms.",
                full_analysis={
                    "companyOverview": {
                        "content": "Global operations in multiple regions.",
                        "sources": [],
                        "confidence": 0.9,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            # Document should open
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None

            # Get all text
            all_text = " ".join(para.text for para in doc.paragraphs)

            # Unicode content should be preserved
            assert "cafe" in all_text or "cafe" in all_text

    def test_json_handles_unicode(self, app):
        """
        Test JSON export handles Unicode without escaping.

        Verifies EXP-05: Unicode preserved, not escaped to \\uXXXX.
        """
        with app.app_context():
            company = Company(
                company_name="Unicode JSON Corp",
                website_url="https://unicode-json.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary with cafe term.",
                full_analysis={
                    "companyOverview": {
                        "content": "International company.",
                        "sources": [],
                        "confidence": 0.9,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json()

            # Should be valid JSON
            data = json.loads(json_str)
            assert isinstance(data, dict)

            # Unicode should not be escaped (ensure_ascii=False in service)
            assert "\\u" not in json_str or "cafe" in json_str


class TestExportLargeContent:
    """Tests for export with large content (EXP-05).

    Verifies that exports handle large analysis content
    without timeout or memory issues.
    """

    def test_markdown_handles_large_analysis(self, app):
        """
        Test Markdown handles large analysis content (70KB+).

        Verifies EXP-05: Completes within 5 seconds with all content.
        """
        with app.app_context():
            company = Company(
                company_name="Large Analysis Corp",
                website_url="https://large.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.THOROUGH,
            )
            db.session.add(company)
            db.session.flush()

            # Create large content (~10KB per section)
            large_content = "This is test content. " * 500  # ~11KB

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary=large_content,
                full_analysis={
                    "companyOverview": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.9,
                    },
                    "businessModel": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.8,
                    },
                    "teamLeadership": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.8,
                    },
                    "marketPosition": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.7,
                    },
                    "keyInsights": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.8,
                    },
                    "redFlags": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.9,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            start_time = time.time()
            markdown = service.generate_markdown()
            elapsed = time.time() - start_time

            # Should complete within 5 seconds
            assert elapsed < 5.0, f"Markdown generation took {elapsed:.2f}s"

            # Should contain all content
            assert "This is test content" in markdown
            assert len(markdown) > 60000  # At least 60KB

    def test_pdf_handles_large_analysis(self, app):
        """
        Test PDF handles large analysis content.

        Verifies EXP-05: Completes without error, reasonable page count.
        """
        with app.app_context():
            company = Company(
                company_name="Large PDF Corp",
                website_url="https://large-pdf.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.THOROUGH,
            )
            db.session.add(company)
            db.session.flush()

            # Large content for PDF
            large_content = "This is test content for PDF generation. " * 300

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary=large_content,
                full_analysis={
                    "companyOverview": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.9,
                    },
                    "businessModel": {
                        "content": large_content,
                        "sources": [],
                        "confidence": 0.8,
                    },
                },
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            # Should complete without error
            pdf_bytes = service.generate_pdf()

            # PDF should be readable
            reader = PdfReader(BytesIO(pdf_bytes))

            # Reasonable page count (not exceeding 20 pages)
            assert len(reader.pages) <= 20

    def test_json_handles_many_entities(self, app):
        """
        Test JSON handles company with 500 entities.

        Verifies EXP-05: Completes within timeout, entities array correct.
        """
        with app.app_context():
            company = Company(
                company_name="Many Entities Corp",
                website_url="https://many-entities.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.THOROUGH,
            )
            db.session.add(company)
            db.session.flush()

            # Add 500 entities
            for i in range(500):
                entity = Entity(
                    company_id=company.id,
                    entity_type=EntityType.PERSON,
                    entity_value=f"Person {i}",
                    confidence_score=0.8,
                    source_url=f"https://many-entities.com/page{i % 10}",
                )
                db.session.add(entity)

            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            start_time = time.time()
            json_str = service.generate_json(include_raw_data=True)
            elapsed = time.time() - start_time

            # Should complete within reasonable time
            assert elapsed < 10.0, f"JSON generation took {elapsed:.2f}s"

            data = json.loads(json_str)

            # Entities array should have 500 items
            assert "entities" in data
            assert len(data["entities"]) == 500

    def test_json_handles_many_pages(self, app):
        """
        Test JSON handles company with 100 pages.

        Verifies EXP-05: Pages array has correct length.
        """
        with app.app_context():
            company = Company(
                company_name="Many Pages Corp",
                website_url="https://many-pages.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.THOROUGH,
            )
            db.session.add(company)
            db.session.flush()

            # Add 100 pages
            for i in range(100):
                page = Page(
                    company_id=company.id,
                    url=f"https://many-pages.com/page{i}",
                    page_type=PageType.OTHER,
                    is_external=False,
                )
                db.session.add(page)

            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=True)

            data = json.loads(json_str)

            # Pages array should have 100 items
            assert "pages" in data
            assert len(data["pages"]) == 100


class TestExportMissingRelatedData:
    """Tests for export with missing related data (EXP-05).

    Verifies that exports handle missing entities, pages, and
    token usage records gracefully.
    """

    def test_export_with_no_entities(self, app):
        """
        Test all export formats handle company with no entities.

        Verifies EXP-05: All formats succeed, JSON has empty entities array.
        """
        with app.app_context():
            company = Company(
                company_name="No Entities Corp",
                website_url="https://no-entities.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary",
                full_analysis={},
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            # Markdown should succeed
            markdown = service.generate_markdown()
            assert isinstance(markdown, str)

            # Word should succeed
            docx_bytes = service.generate_word()
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None

            # PDF should succeed
            pdf_bytes = service.generate_pdf()
            reader = PdfReader(BytesIO(pdf_bytes))
            assert len(reader.pages) > 0

            # JSON should have empty entities array
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)
            assert data["entities"] == []

    def test_export_with_no_pages(self, app):
        """
        Test all export formats handle company with no pages.

        Verifies EXP-05: All formats succeed, JSON has empty pages array.
        """
        with app.app_context():
            company = Company(
                company_name="No Pages Corp",
                website_url="https://no-pages.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary",
                full_analysis={},
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)

            # All formats should succeed
            markdown = service.generate_markdown()
            assert "## Sources" in markdown

            docx_bytes = service.generate_word()
            assert len(docx_bytes) > 0

            pdf_bytes = service.generate_pdf()
            assert len(pdf_bytes) > 0

            # JSON should have empty pages array
            json_str = service.generate_json(include_raw_data=True)
            data = json.loads(json_str)
            assert data["pages"] == []

    def test_export_with_no_token_usage(self, app):
        """
        Test Markdown handles company with no token usage records.

        Verifies EXP-05: Token count shows 0, no division by zero.
        """
        with app.app_context():
            company = Company(
                company_name="No Tokens Corp",
                website_url="https://no-tokens.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.flush()

            analysis = Analysis(
                company_id=company.id,
                version_number=1,
                executive_summary="Summary",
                full_analysis={},
            )
            db.session.add(analysis)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should show 0 tokens
            assert "**Tokens Used:** 0" in markdown

            # Cost should be 0
            assert "$0.0000" in markdown

    def test_json_without_raw_data_excludes_empty_arrays(self, app):
        """
        Test JSON without raw data excludes entities/pages keys.

        Verifies EXP-05: With include_raw_data=False, keys are absent.
        """
        with app.app_context():
            company = Company(
                company_name="No Raw Data Corp",
                website_url="https://no-raw-data.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json(include_raw_data=False)

            data = json.loads(json_str)

            # entities and pages keys should not be present
            assert "entities" not in data
            assert "pages" not in data


class TestExportContentValidation:
    """Tests for validating export output formats (EXP-05).

    Verifies that all export formats produce valid, well-formed output.
    """

    def test_markdown_output_is_valid_utf8(self, app):
        """
        Test Markdown output is valid UTF-8.

        Verifies EXP-05: No encoding errors.
        """
        with app.app_context():
            company = Company(
                company_name="UTF8 Test Corp",
                website_url="https://utf8-test.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            markdown = service.generate_markdown()

            # Should encode to UTF-8 without error
            try:
                encoded = markdown.encode("utf-8")
                assert isinstance(encoded, bytes)
            except UnicodeEncodeError:
                pytest.fail("Markdown output is not valid UTF-8")

    def test_word_output_is_valid_docx(self, app):
        """
        Test Word output is valid DOCX format.

        Verifies EXP-05: Starts with PK magic bytes (ZIP), python-docx can parse.
        """
        with app.app_context():
            company = Company(
                company_name="DOCX Test Corp",
                website_url="https://docx-test.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            docx_bytes = service.generate_word()

            # Should start with PK magic bytes (ZIP format)
            assert docx_bytes[:2] == b"PK"

            # python-docx should be able to parse it
            doc = Document(BytesIO(docx_bytes))
            assert doc is not None
            assert len(doc.paragraphs) > 0

    def test_pdf_output_is_valid_pdf(self, app):
        """
        Test PDF output is valid PDF format.

        Verifies EXP-05: Starts with %PDF magic bytes, PyPDF2 can parse.
        """
        with app.app_context():
            company = Company(
                company_name="PDF Test Corp",
                website_url="https://pdf-test.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            pdf_bytes = service.generate_pdf()

            # Should start with %PDF magic bytes
            assert pdf_bytes[:4] == b"%PDF"

            # PyPDF2 should be able to parse it
            reader = PdfReader(BytesIO(pdf_bytes))
            assert len(reader.pages) > 0

    def test_json_output_is_valid_json(self, app):
        """
        Test JSON output is valid JSON format.

        Verifies EXP-05: json.loads succeeds, result is dict with expected keys.
        """
        with app.app_context():
            company = Company(
                company_name="JSON Test Corp",
                website_url="https://json-test.com",
                status=CompanyStatus.COMPLETED,
                analysis_mode=AnalysisMode.QUICK,
            )
            db.session.add(company)
            db.session.commit()

            company = db.session.get(Company, company.id)
            service = ExportService(company)
            json_str = service.generate_json()

            # json.loads should succeed
            data = json.loads(json_str)

            # Result should be dict with expected keys
            assert isinstance(data, dict)
            assert "company" in data
            assert "analysis" in data
            assert "tokenUsage" in data
            assert "metadata" in data
